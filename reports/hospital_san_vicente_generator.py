"""
Generador de reportes mejorado para Hospital San Vicente
Integra la plantilla específica con variables dinámicas y constantes fijas
Basado en GUIA_PRACTICA_ELABORACION_INFORMES_MENSUALES y PLANTILLA_BASE_INFORME_TECNICO
"""

import yaml
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from io import BytesIO
import matplotlib.pyplot as plt
from typing import Dict, Any, Tuple, Optional
import os
import calendar

# Importar módulos existentes del sistema
from visualisations.Preventivos import (
    generate_order_area_plot,
    generate_plagas_timeseries_facet,
    generate_total_plagas_trend_plot
)
from visualisations.Roedores import (
    generate_roedores_station_status_plot,
    plot_tendencia_eliminacion_mensual
)
from visualisations.Lamparas import (
    plot_estado_lamparas_por_mes,
    plot_estado_lamparas_con_leyenda,
    plot_capturas_especies_por_mes,
    plot_tendencia_total_capturas
)

class HospitalSanVicenteReportGenerator:
    """
    Generador de reportes específico para Hospital San Vicente
    Implementa la plantilla oficial con variables dinámicas automatizadas
    """
    
    def __init__(self):
        self.template_config = self._load_template_config()
        self.current_date = datetime.now()
        
    def _load_template_config(self) -> Dict[str, Any]:
        """Carga la configuración de la plantilla desde el archivo YAML."""
        config_path = "config/hospital_san_vicente_template.yaml"
        try:
            with open(config_path, 'r', encoding='utf-8') as file:
                return yaml.safe_load(file)
        except FileNotFoundError:
            raise FileNotFoundError(f"No se encontró la plantilla en {config_path}")
        except yaml.YAMLError as e:
            raise ValueError(f"Error al cargar la plantilla YAML: {e}")
    
    def calculate_dynamic_variables(self, df_preventivo: pd.DataFrame, 
                                  df_roedores: pd.DataFrame, 
                                  df_lamparas: pd.DataFrame, 
                                  sede: str) -> Dict[str, Any]:
        """
        Calcula todas las variables dinámicas basadas en los datos reales
        Según la GUIA_PRACTICA_ELABORACION_INFORMES_MENSUALES
        """
        # Obtener configuración de la sede
        sede_config = self.template_config['sedes'][sede]
        
        # Filtrar datos por sede
        df_preventivo_filtered = df_preventivo[df_preventivo['Sede'] == sede]
        df_roedores_filtered = df_roedores[df_roedores['Sede'] == sede]
        df_lamparas_filtered = df_lamparas[df_lamparas['Sede'] == sede]
        
        # Variables temporales
        mes_actual = self.current_date.month
        año_actual = self.current_date.year
        ultimo_dia_mes = calendar.monthrange(año_actual, mes_actual)[1]
        fecha_elaboracion = f"{ultimo_dia_mes:02d}/{mes_actual:02d}/{año_actual}"
        
        # Calcular métricas principales de preventivos
        total_ordenes = len(df_preventivo_filtered) if not df_preventivo_filtered.empty else 0
        areas_con_plagas = 0
        total_plagas = 0
        
        if not df_preventivo_filtered.empty:
            # Contar áreas con evidencia de plagas
            plagas_columns = ['Cucarachas Alemanas', 'Cucarachas Americanas', 'Moscas', 'Zancudos', 'Ratas', 'Ratones', 'Hormigas']
            for col in plagas_columns:
                if col in df_preventivo_filtered.columns:
                    areas_con_plagas += (df_preventivo_filtered[col] > 0).sum()
                    total_plagas += df_preventivo_filtered[col].sum()
        
        # Determinar grado de infestación
        grado_infestacion = self._determinar_grado_infestacion(total_plagas)
        
        # Calcular porcentaje de cumplimiento basado en el grado de infestación
        porcentaje_cumplimiento = self._calcular_porcentaje_cumplimiento(grado_infestacion, total_plagas)
        
        # Análisis de roedores
        estaciones_activas = 0
        consumo_total = 0
        if not df_roedores_filtered.empty:
            estaciones_activas = len(df_roedores_filtered)
            if 'Consumo' in df_roedores_filtered.columns:
                consumo_total = df_roedores_filtered['Consumo'].sum()
        
        # Análisis de lámparas  
        lamparas_funcionales = 0
        lamparas_saturadas = 0
        total_capturas = 0
        if not df_lamparas_filtered.empty:
            lamparas_funcionales = len(df_lamparas_filtered)
            if 'Estado' in df_lamparas_filtered.columns:
                lamparas_saturadas = (df_lamparas_filtered['Estado'] == 'Saturada').sum()
            if 'Capturas' in df_lamparas_filtered.columns:
                total_capturas = df_lamparas_filtered['Capturas'].sum()
        
        # Número de mes del ciclo (asumimos que empezó en enero)
        numero_mes_ciclo = mes_actual
        
        return {
            # Variables temporales
            'fecha_elaboracion': fecha_elaboracion,
            'mes_nombre': self.template_config['configuracion_temporal']['meses_español'][mes_actual],
            'año': año_actual,
            'numero_mes_texto': self.template_config['configuracion_temporal']['numeros_ordinales'][numero_mes_ciclo],
            'numero_mes_ciclo': numero_mes_ciclo,
            
            # Variables de sede
            'sede': sede,
            'nombre_completo_sede': sede_config['nombre_completo'],
            'codigo_sede': sede_config['codigo_sede'],
            'direccion': sede_config['direccion'],
            'telefono': sede_config['telefono'],
            'numero_bloques': sede_config['numero_bloques'],
            
            # Métricas principales (dinámicas)
            'ordenes_solicitadas': total_ordenes,
            'ordenes_realizadas': total_ordenes,  # Usualmente igual según guía
            'porcentaje_cumplimiento': porcentaje_cumplimiento,
            'total_areas_controladas': len(df_preventivo_filtered) if not df_preventivo_filtered.empty else 0,
            'areas_con_plagas': areas_con_plagas,
            'total_plagas': total_plagas,
            'grado_infestacion': grado_infestacion,
            
            # Métricas de roedores
            'numero_estaciones_roedores': estaciones_activas,
            'consumo_total_roedores': consumo_total,
            
            # Métricas de lámparas
            'numero_lamparas': lamparas_funcionales,
            'lamparas_saturadas': lamparas_saturadas,
            'total_capturas_insectos': total_capturas,
            
            # Análisis por especies
            'tiene_cucarachas': self._tiene_especie(df_preventivo_filtered, ['Cucarachas Alemanas', 'Cucarachas Americanas']),
            'tiene_voladores': self._tiene_especie(df_preventivo_filtered, ['Moscas', 'Zancudos']),
            'tiene_roedores': self._tiene_especie(df_preventivo_filtered, ['Ratas', 'Ratones']),
            'tiene_hormigas': self._tiene_especie(df_preventivo_filtered, ['Hormigas']),
            
            # Páginas del informe
            'total_paginas': sede_config['paginas_informe']
        }
    
    def _determinar_grado_infestacion(self, total_plagas: int) -> str:
        """Determina el grado de infestación según los rangos de la guía práctica."""
        grados = self.template_config['grados_infestacion']
        
        if total_plagas <= grados['bajo']['limite']:
            return 'bajo'
        elif total_plagas <= grados['medio']['limite']:
            return 'medio'
        else:
            return 'alto'
    
    def _calcular_porcentaje_cumplimiento(self, grado_infestacion: str, total_plagas: int) -> int:
        """Calcula el porcentaje de cumplimiento basado en el rendimiento."""
        if grado_infestacion == 'bajo' and total_plagas == 0:
            return 83  # Excelente rendimiento
        elif grado_infestacion == 'bajo':
            return 80  # Buen rendimiento
        elif grado_infestacion == 'medio':
            return 75  # Rendimiento con incidencias menores
        else:
            return 70  # Rendimiento por debajo del estándar
    
    def _tiene_especie(self, df: pd.DataFrame, columnas: list) -> bool:
        """Verifica si hay evidencia de una especie específica."""
        if df.empty:
            return False
        
        for col in columnas:
            if col in df.columns and df[col].sum() > 0:
                return True
        return False
    
    def _format_template_text(self, template: str, variables: Dict[str, Any]) -> str:
        """Formatea un template con las variables proporcionadas."""
        try:
            return template.format(**variables)
        except KeyError as e:
            print(f"⚠️  Variable faltante en template: {e}")
            return template
    
    def _add_document_header(self, doc: Document, variables: Dict[str, Any]):
        """Añade el encabezado oficial del documento."""
        # Logo
        if os.path.exists('Logo/logo2021.png'):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture('Logo/logo2021.png', width=Inches(3.25))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del documento (FIJO)
        doc_config = self.template_config['document_config']
        
        # Título principal
        title = doc.add_heading(f"INFORME TÉCNICO FINAL", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Código y versión
        info_para = doc.add_paragraph()
        info_para.add_run(f"Código: {doc_config['codigo_formato']} | ")
        info_para.add_run(f"Versión: {doc_config['version']} | ")
        info_para.add_run(f"Fecha de formato: {doc_config['fecha_formato']}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espacio
        
        # Información del cliente
        client_heading = doc.add_heading("INFORMACIÓN DEL CLIENTE", level=2)
        
        # Tabla de información del cliente
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Llenar tabla con datos variables
        client_data = [
            ("CLIENTE:", variables['nombre_completo_sede']),
            ("FECHA DE ELABORACIÓN:", variables['fecha_elaboracion']),
            ("DIRECCIÓN:", variables['direccion']),
            ("SEDE:", variables['codigo_sede']),
            ("TELÉFONO:", variables['telefono']),
            ("SECTOR:", doc_config['sector']),
            ("MES DE ANÁLISIS:", variables['mes_nombre']),
            ("AÑO DE ANÁLISIS:", str(variables['año']))
        ]
        
        for i, (key, value) in enumerate(client_data):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
            # Hacer la primera columna en negrita
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    def _add_metrics_section(self, doc: Document, variables: Dict[str, Any]):
        """Añade la sección de métricas principales."""
        doc.add_heading("MÉTRICAS DEL PERÍODO", level=2)
        
        # Tabla de métricas
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        metrics_data = [
            ("ÓRDENES SOLICITADAS:", f"{variables['ordenes_solicitadas']} PLAN MTO PREVENTIVO"),
            ("ÓRDENES REALIZADAS:", f"{variables['ordenes_realizadas']} PLAN MTO PREVENTIVO"),
            ("PORCENTAJE CUMPLIMIENTO:", f"{variables['porcentaje_cumplimiento']}%"),
            ("BLOQUES ASISTENCIALES:", str(variables['numero_bloques'])),
            ("GRADO DE INFESTACIÓN:", variables['grado_infestacion'].upper())
        ]
        
        for i, (key, value) in enumerate(metrics_data):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            
            # Colorear el grado de infestación
            if "GRADO DE INFESTACIÓN" in key:
                color_map = {'BAJO': RGBColor(0, 128, 0), 'MEDIO': RGBColor(255, 165, 0), 'ALTO': RGBColor(255, 0, 0)}
                if variables['grado_infestacion'].upper() in color_map:
                    table.cell(i, 1).paragraphs[0].runs[0].font.color.rgb = color_map[variables['grado_infestacion'].upper()]
    
    def _add_program_description(self, doc: Document, variables: Dict[str, Any]):
        """Añade la descripción del programa con texto dinámico."""
        doc.add_heading("DESCRIPCIÓN DEL PROGRAMA", level=2)
        
        # Contenido fijo
        fixed_content = self.template_config['contenido_fijo']['descripcion_programa']['contenido']
        doc.add_paragraph(fixed_content)
        
        # Contenido dinámico basado en plantillas
        intro_template = self.template_config['plantillas_texto']['introduccion_ciclo']['template']
        intro_text = self._format_template_text(intro_template, variables)
        doc.add_paragraph(intro_text)
        
        # Métricas principales con template dinámico
        metricas_template = self.template_config['plantillas_texto']['metricas_principales']['template']
        metricas_text = self._format_template_text(metricas_template, variables)
        doc.add_paragraph(metricas_text)
        
        # Nota estándar (fijo)
        nota_content = self.template_config['contenido_fijo']['nota_consolidados']['contenido']
        doc.add_paragraph(nota_content)
    
    def _add_analysis_sections(self, doc: Document, df_preventivo: pd.DataFrame, 
                             df_roedores: pd.DataFrame, df_lamparas: pd.DataFrame, 
                             variables: Dict[str, Any]):
        """Añade las secciones de análisis con gráficas y tablas."""
        
        # SECCIÓN 1: CONTROL DE INSECTOS RASTREROS
        doc.add_heading("1. REGISTRO CONTROL DE INSECTOS RASTREROS - PREVENTIVOS", level=2)
        
        # Análisis basado en datos reales
        if variables['total_plagas'] == 0:
            # Usar plantilla para caso sin plagas
            sin_plagas_template = self.template_config['plantillas_texto']['analisis_sin_plagas']['template']
            sin_plagas_text = self._format_template_text(sin_plagas_template, variables)
            doc.add_paragraph(sin_plagas_text)
            
            # Tendencia positiva
            tendencia_template = self.template_config['plantillas_texto']['tendencia_positiva']['template']
            tendencia_text = self._format_template_text(tendencia_template, variables)
            doc.add_paragraph(tendencia_text)
        else:
            # Análisis detallado cuando hay plagas
            doc.add_paragraph(f"Durante el período se registraron un total de {variables['total_plagas']} "
                            f"especímenes de plagas distribuidos en {variables['areas_con_plagas']} áreas.")
        
        # Gráficas de preventivos
        if not df_preventivo.empty:
            self._add_preventivos_charts(doc, df_preventivo)
        
        # SECCIÓN 2: CONTROL DE ROEDORES  
        doc.add_heading("2. REGISTRO CONTROL DE ROEDORES - CONSOLIDADO MENSUAL", level=2)
        
        doc.add_paragraph(f"Se monitorearon {variables['numero_estaciones_roedores']} estaciones portacebos "
                         f"con un consumo total de {variables['consumo_total_roedores']} gramos de rodenticida.")
        
        # Gráficas de roedores
        if not df_roedores.empty:
            self._add_roedores_charts(doc, df_roedores)
        
        # SECCIÓN 3: CONTROL DE INSECTOS VOLADORES
        doc.add_heading("3. REGISTRO CONTROL DE INSECTOS VOLADORES - LÁMPARAS", level=2)
        
        doc.add_paragraph(f"Se mantuvieron {variables['numero_lamparas']} lámparas de control, "
                         f"registrándose {variables['lamparas_saturadas']} lámparas saturadas que requieren mantenimiento.")
        
        # Gráficas de lámparas
        if not df_lamparas.empty:
            self._add_lamparas_charts(doc, df_lamparas)
    
    def _add_preventivos_charts(self, doc: Document, df_preventivo: pd.DataFrame):
        """Añade gráficas de la sección preventivos."""
        try:
            # Gráfica 1: Órdenes vs Áreas
            table_data, figure = generate_order_area_plot(df_preventivo)
            if figure:
                doc.add_paragraph("Gráfica #1: Cantidad de órdenes vs cantidad de áreas", style='Intense Quote')
                self._add_plot_to_doc(doc, figure)
                self._add_table_to_doc(doc, table_data)
            
            # Gráfica 2: Especies de plagas
            table_data, figure = generate_plagas_timeseries_facet(df_preventivo)
            if figure:
                doc.add_paragraph("Gráfica #2: Relación por especie encontrada", style='Intense Quote')
                self._add_plot_to_doc(doc, figure)
                self._add_table_to_doc(doc, table_data)
            
            # Gráfica 3: Tendencia total
            table_data, figure = generate_total_plagas_trend_plot(df_preventivo)
            if figure:
                doc.add_paragraph("Gráfica #3: Tendencia de eliminación mensual", style='Intense Quote')
                self._add_plot_to_doc(doc, figure)
                self._add_table_to_doc(doc, table_data)
                
        except Exception as e:
            print(f"⚠️  Error generando gráficas de preventivos: {e}")
    
    def _add_roedores_charts(self, doc: Document, df_roedores: pd.DataFrame):
        """Añade gráficas de la sección roedores."""
        try:
            # Estado de estaciones
            table_data, figure = generate_roedores_station_status_plot(df_roedores)
            if figure:
                doc.add_paragraph("Estado de las estaciones portacebos", style='Intense Quote')
                self._add_plot_to_doc(doc, figure)
                self._add_table_to_doc(doc, table_data)
            
            # Tendencia de eliminación
            table_data, figure = plot_tendencia_eliminacion_mensual(df_roedores)
            if figure:
                doc.add_paragraph("Tendencia de consumo mensual", style='Intense Quote')
                self._add_plot_to_doc(doc, figure)
                self._add_table_to_doc(doc, table_data)
                
        except Exception as e:
            print(f"⚠️  Error generando gráficas de roedores: {e}")
    
    def _add_lamparas_charts(self, doc: Document, df_lamparas: pd.DataFrame):
        """Añade gráficas de la sección lámparas."""
        try:
            # Estado mensual
            table_data, figure = plot_estado_lamparas_por_mes(df_lamparas)
            if figure:
                doc.add_paragraph("Estado de las lámparas por mes", style='Intense Quote')
                self._add_plot_to_doc(doc, figure)
                self._add_table_to_doc(doc, table_data)
            
            # Estado con leyenda
            table_data, figure = plot_estado_lamparas_con_leyenda(df_lamparas)
            if figure:
                doc.add_paragraph("Estado detallado de las lámparas", style='Intense Quote')
                self._add_plot_to_doc(doc, figure)
                self._add_table_to_doc(doc, table_data)
            
            # Capturas por especies
            table_data, figure = plot_capturas_especies_por_mes(df_lamparas)
            if figure:
                doc.add_paragraph("Capturas por especies", style='Intense Quote')
                self._add_plot_to_doc(doc, figure)
                self._add_table_to_doc(doc, table_data)
            
            # Tendencia de capturas
            table_data, figure = plot_tendencia_total_capturas(df_lamparas)
            if figure:
                doc.add_paragraph("Tendencia total de capturas", style='Intense Quote')
                self._add_plot_to_doc(doc, figure)
                self._add_table_to_doc(doc, table_data)
                
        except Exception as e:
            print(f"⚠️  Error generando gráficas de lámparas: {e}")
    
    def _add_plot_to_doc(self, doc: Document, fig):
        """Añade una gráfica de matplotlib al documento."""
        if fig is None:
            return False
        
        try:
            with BytesIO() as image_stream:
                fig.savefig(image_stream, format='png', bbox_inches='tight', dpi=300)
                image_stream.seek(0)
                doc.add_picture(image_stream, width=Inches(5.5))
            plt.close(fig)
            return True
        except Exception as e:
            print(f"Error añadiendo gráfica: {e}")
            if fig is not None:
                plt.close(fig)
            return False
    
    def _add_table_to_doc(self, doc: Document, table_data: pd.DataFrame):
        """Añade una tabla de datos al documento."""
        if table_data is None or table_data.empty:
            return
        
        try:
            # Crear tabla en Word
            table = doc.add_table(rows=len(table_data) + 1, cols=len(table_data.columns))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Añadir encabezados
            for i, column in enumerate(table_data.columns):
                cell = table.cell(0, i)
                cell.text = str(column)
                cell.paragraphs[0].runs[0].bold = True
            
            # Añadir datos
            for i, row in table_data.iterrows():
                for j, value in enumerate(row):
                    table.cell(i + 1, j).text = str(value)
                    
        except Exception as e:
            print(f"Error añadiendo tabla: {e}")
    
    def _add_recommendations(self, doc: Document, variables: Dict[str, Any]):
        """Añade la sección de recomendaciones."""
        doc.add_heading("RECOMENDACIONES", level=2)
        
        # Recomendaciones generales (fijas)
        recom_config = self.template_config['contenido_fijo']['recomendaciones_generales']
        doc.add_heading("Recomendaciones Generales:", level=3)
        
        for i, recom in enumerate(recom_config['lista'], 1):
            para = doc.add_paragraph()
            para.add_run(f"{i}. {recom['texto']}").bold = False
            para.add_run(f"\n   Responsable: {recom['responsable']}").italic = True
        
        # Recomendaciones específicas por sede
        sede_config = self.template_config['sedes'][variables['sede']]
        if 'recomendaciones_especificas' in sede_config:
            doc.add_heading(f"Recomendaciones Específicas para {variables['sede']}:", level=3)
            
            for i, recom in enumerate(sede_config['recomendaciones_especificas'], 1):
                doc.add_paragraph(f"{i}. {recom}")
    
    def _add_footer_contact(self, doc: Document):
        """Añade la información de contacto al final."""
        doc.add_paragraph()  # Espacio
        
        # Despedida
        para = doc.add_paragraph("Cordialmente,")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph()  # Espacio para firma
        doc.add_paragraph()  # Espacio para firma
        
        # Información de contacto
        contacto = self.template_config['contacto']
        
        # Nombre y cargo
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(contacto['nombre'])
        name_run.bold = True
        name_run.font.size = Pt(12)
        
        cargo_para = doc.add_paragraph()
        cargo_run = cargo_para.add_run(contacto['cargo'])
        cargo_run.bold = True
        cargo_run.font.size = Pt(10)
        
        # Datos de contacto
        contact_para = doc.add_paragraph()
        contact_para.add_run(f"Teléfonos: {contacto['telefono']} - ")
        contact_para.add_run(f"Celular: {contacto['celular']}\n")
        contact_para.add_run(f"e-mail: {contacto['email']}")
    
    def generate_complete_report(self, df_preventivo: pd.DataFrame, 
                               df_roedores: pd.DataFrame, 
                               df_lamparas: pd.DataFrame, 
                               sede: str) -> str:
        """
        Genera un reporte completo usando la nueva plantilla del Hospital San Vicente.
        
        Returns:
            str: Ruta del archivo generado
        """
        print(f"\n🏥 Generando reporte Hospital San Vicente para {sede}")
        print("=" * 60)
        
        # Calcular variables dinámicas
        print("📊 Calculando variables dinámicas...")
        variables = self.calculate_dynamic_variables(df_preventivo, df_roedores, df_lamparas, sede)
        
        # Crear documento
        print("📄 Creando documento...")
        doc = Document()
        
        # Construir el reporte sección por sección
        print("🔧 Construyendo encabezado...")
        self._add_document_header(doc, variables)
        
        print("📈 Añadiendo métricas...")
        self._add_metrics_section(doc, variables)
        
        print("📝 Añadiendo descripción del programa...")
        self._add_program_description(doc, variables)
        
        print("📊 Procesando análisis y gráficas...")
        self._add_analysis_sections(doc, df_preventivo, df_roedores, df_lamparas, variables)
        
        print("💡 Añadiendo recomendaciones...")
        self._add_recommendations(doc, variables)
        
        print("📞 Añadiendo información de contacto...")
        self._add_footer_contact(doc)
        
        # Guardar documento
        output_path = f"outputs/informe_hospital_san_vicente_{sede.lower()}_{variables['año']}_{variables['mes_nombre'].lower()}.docx"
        doc.save(output_path)
        
        print(f"✅ Reporte completado: {output_path}")
        
        # Mostrar resumen
        print(f"\n📋 RESUMEN DEL REPORTE:")
        print(f"   📅 Período: {variables['mes_nombre']} {variables['año']}")
        print(f"   🏢 Sede: {variables['nombre_completo_sede']}")
        print(f"   📊 Órdenes procesadas: {variables['ordenes_solicitadas']}")
        print(f"   🎯 Cumplimiento: {variables['porcentaje_cumplimiento']}%")
        print(f"   🐛 Plagas totales: {variables['total_plagas']}")
        print(f"   📈 Grado infestación: {variables['grado_infestacion'].upper()}")
        
        return output_path


def generate_hospital_san_vicente_report(df_preventivo: pd.DataFrame, 
                                       df_roedores: pd.DataFrame, 
                                       df_lamparas: pd.DataFrame, 
                                       sede: str) -> str:
    """
    Función de conveniencia para generar reportes del Hospital San Vicente.
    
    Args:
        df_preventivo: DataFrame con datos de servicios preventivos
        df_roedores: DataFrame con datos de control de roedores
        df_lamparas: DataFrame con datos de control de lámparas  
        sede: Nombre de la sede ('Rionegro' o 'Medellín')
        
    Returns:
        str: Ruta del archivo de reporte generado
    """
    generator = HospitalSanVicenteReportGenerator()
    return generator.generate_complete_report(df_preventivo, df_roedores, df_lamparas, sede)
