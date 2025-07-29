import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Tuple, List

from visualisations.Preventivos import (
    generate_order_area_plot,
    generate_plagas_timeseries_facet,
    generate_total_plagas_trend_plot)

from visualisations.Roedores import (
    generate_roedores_station_status_plot,
    plot_tendencia_eliminacion_mensual)

from visualisations.Lamparas import (
    plot_estado_lamparas_por_mes,
    plot_estado_lamparas_con_leyenda,
    plot_capturas_especies_por_mes,
    plot_tendencia_total_capturas)

from llm_integration.prompt_generator import LLMPromptGenerator, print_prompt_with_separator
from config.settings import DOCUMENT_CONFIG

class ReportDataManager:
    """Gestiona los datos y tablas para la generación de reportes."""
    
    def __init__(self):
        self.llm_generator = LLMPromptGenerator()
        self.section_data = {
            'preventivos': {},
            'roedores': {},
            'lamparas': {}
        }
        # Almacenar prompts generados para manipulación posterior
        self.generated_prompts = {
            'table_prompts': [],  # Lista de prompts individuales por tabla
            'section_prompts': {},  # Prompts de resumen por sección
            'general_prompt': None  # Prompt de resumen general
        }
    
    def store_table_data(self, section: str, table_name: str, table_data: pd.DataFrame):
        """Almacena los datos de una tabla para su uso posterior."""
        self.section_data[section][table_name] = table_data.copy()
    
    def generate_table_prompt(self, section: str, table_name: str, table_data: pd.DataFrame, sede: str):
        """Genera y almacena el prompt para una tabla específica."""
        # Almacenar los datos
        self.store_table_data(section, table_name, table_data)
        
        # Generar prompt
        prompt = self.llm_generator.generate_table_description_prompt(
            table_data, section, table_name, sede
        )
        
        # Almacenar prompt con metadatos para manipulación posterior
        prompt_data = {
            'id': f"{section}_{table_name}",
            'section': section,
            'table_name': table_name,
            'sede': sede,
            'prompt': prompt,
            'table_data': table_data.copy(),
            'title': f"{section.upper()} - {table_name.replace('_', ' ').title()}"
        }
        
        self.generated_prompts['table_prompts'].append(prompt_data)
        
        return prompt_data
    
    def generate_section_summary_prompt(self, section: str, sede: str):
        """Genera y almacena el prompt de resumen para una sección completa."""
        if section in self.section_data and self.section_data[section]:
            prompt = self.llm_generator.generate_section_summary_prompt(
                self.section_data[section], section, sede
            )
            
            # Almacenar prompt de sección
            prompt_data = {
                'section': section,
                'sede': sede,
                'prompt': prompt,
                'title': f"RESUMEN SECCIÓN - {section.upper()}",
                'section_tables': self.section_data[section].copy()
            }
            
            self.generated_prompts['section_prompts'][section] = prompt_data
            
            return prompt_data
        return None
    
    def generate_general_summary_prompt(self, sede: str):
        """Genera y almacena el prompt de resumen general."""
        prompt = self.llm_generator.generate_general_summary_prompt(
            self.section_data, sede
        )
        
        # Almacenar prompt general
        prompt_data = {
            'sede': sede,
            'prompt': prompt,
            'title': "RESUMEN GENERAL DEL REPORTE",
            'all_data': self.section_data.copy()
        }
        
        self.generated_prompts['general_prompt'] = prompt_data
        
        return prompt_data
    
    def get_all_prompts(self):
        """Retorna todos los prompts generados para manipulación."""
        return self.generated_prompts
    
    def get_table_prompts(self):
        """Retorna solo los prompts de tablas individuales."""
        return self.generated_prompts['table_prompts']
    
    def get_section_prompts(self):
        """Retorna solo los prompts de resúmenes de sección."""
        return self.generated_prompts['section_prompts']
    
    def get_general_prompt(self):
        """Retorna el prompt de resumen general."""
        return self.generated_prompts['general_prompt']
    
    def update_prompt(self, prompt_id: str, new_prompt: str):
        """Permite modificar un prompt específico."""
        # Buscar en prompts de tabla
        for prompt_data in self.generated_prompts['table_prompts']:
            if prompt_data['id'] == prompt_id:
                prompt_data['prompt'] = new_prompt
                return True
        
        # Buscar en prompts de sección
        for section, prompt_data in self.generated_prompts['section_prompts'].items():
            if f"section_{section}" == prompt_id:
                prompt_data['prompt'] = new_prompt
                return True
        
        # Verificar prompt general
        if prompt_id == "general" and self.generated_prompts['general_prompt']:
            self.generated_prompts['general_prompt']['prompt'] = new_prompt
            return True
        
        return False
    
    def export_prompts_to_file(self, filename: str = "generated_prompts.txt"):
        """Exporta todos los prompts a un archivo de texto para fácil manipulación."""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("PROMPTS GENERADOS PARA LLM - SISTEMA SERVIPLAGAS\n")
            f.write("=" * 80 + "\n\n")
            
            # Prompts de tablas individuales
            f.write("🔍 PROMPTS DE ANÁLISIS DE TABLAS INDIVIDUALES\n")
            f.write("-" * 50 + "\n\n")
            
            for i, prompt_data in enumerate(self.generated_prompts['table_prompts'], 1):
                f.write(f"📊 PROMPT {i}: {prompt_data['title']}\n")
                f.write("=" * 60 + "\n")
                f.write(prompt_data['prompt'])
                f.write("\n\n" + "="*60 + "\n\n")
            
            # Prompts de sección
            f.write("📋 PROMPTS DE RESUMEN POR SECCIÓN\n")
            f.write("-" * 50 + "\n\n")
            
            for section, prompt_data in self.generated_prompts['section_prompts'].items():
                f.write(f"📂 {prompt_data['title']}\n")
                f.write("=" * 60 + "\n")
                f.write(prompt_data['prompt'])
                f.write("\n\n" + "="*60 + "\n\n")
            
            # Prompt general
            if self.generated_prompts['general_prompt']:
                f.write("🌐 PROMPT DE RESUMEN GENERAL\n")
                f.write("-" * 50 + "\n\n")
                f.write(f"📑 {self.generated_prompts['general_prompt']['title']}\n")
                f.write("=" * 60 + "\n")
                f.write(self.generated_prompts['general_prompt']['prompt'])
                f.write("\n\n" + "="*60 + "\n\n")
        
        print(f"✅ Prompts exportados a: {filename}")
        return filename


def _add_prompt_to_doc(doc: Document, prompt_text: str, title: str):
    """Añade un prompt del LLM al documento para revisión."""
    # Solo añadir prompts si está habilitado en la configuración
    if not DOCUMENT_CONFIG.get('include_prompts_in_document', False):
        return
    
    # Añadir separador visual
    doc.add_paragraph()
    
    # Título del prompt
    prompt_heading = doc.add_heading(f"🤖 PROMPT LLM: {title}", level=4)
    prompt_heading.style.font.color.rgb = RGBColor(0, 102, 204)  # Azul
    
    # Contenido del prompt en un cuadro de texto
    prompt_para = doc.add_paragraph()
    prompt_para.style = 'Intense Quote'
    
    # Dividir el prompt en líneas para mejor formato
    if DOCUMENT_CONFIG.get('prompt_style', 'detailed') == 'compact':
        # Versión compacta: solo primeras 3 líneas de instrucciones
        lines = prompt_text.strip().split('\n')[:10]
        prompt_para.add_run("PROMPT (versión compacta):\n").font.bold = True
    else:
        # Versión completa
        lines = prompt_text.strip().split('\n')
    
    for i, line in enumerate(lines):
        if line.strip():
            if i > 0:
                prompt_para.add_run('\n')
            run = prompt_para.add_run(line.strip())
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)  # Gris oscuro
    
    # Placeholder para la respuesta solo si está habilitado
    if DOCUMENT_CONFIG.get('show_llm_placeholders', True):
        response_heading = doc.add_heading("📝 RESPUESTA DEL LLM:", level=4)
        response_heading.style.font.color.rgb = RGBColor(0, 153, 76)  # Verde
        
        response_para = doc.add_paragraph()
        response_para.style = 'Normal'
        response_run = response_para.add_run("[AQUÍ IRÁ LA RESPUESTA DEL LLM BASADA EN EL PROMPT ANTERIOR]")
        response_run.font.italic = True
        response_run.font.color.rgb = RGBColor(128, 128, 128)  # Gris claro
    
    # Separador
    doc.add_paragraph()


# Utilidad para agregar un gráfico de matplotlib directamente al doc
def add_plot_to_doc(doc, fig):
    """Add a matplotlib figure to a Word document"""
    if fig is None:
        print("Warning: Figure is None, skipping plot addition")
        return False

    try:
        with BytesIO() as image_stream:
            fig.savefig(image_stream, format='png', bbox_inches='tight', dpi=300)
            image_stream.seek(0)
            doc.add_picture(image_stream, width=Inches(5.5))
        plt.close(fig)
        return True
    except Exception as e:
        print(f"Error adding plot to document: {e}")
        if fig is not None:
            plt.close(fig)
        return False

# Función que genera el reporte completo con integración LLM
def generate_enhanced_report(df_preventivo: pd.DataFrame, 
                           df_roedores: pd.DataFrame, 
                           df_lamparas: pd.DataFrame, 
                           sede: str) -> str:
    """
    Genera un reporte completo con análisis de datos usando LLM.
    
    Args:
        df_preventivo: DataFrame con datos de servicios preventivos
        df_roedores: DataFrame con datos de control de roedores  
        df_lamparas: DataFrame con datos de control de lámparas
        sede: Nombre de la sede para filtrar datos
        
    Returns:
        str: Ruta del archivo de reporte generado
    """
    print(f"\n🚀 Iniciando generación de reporte para sede: {sede}")
    print("=" * 60)
    
    # Filtrar datos por sede
    df_preventivo_filtered = df_preventivo[df_preventivo['Sede'] == sede].copy()
    df_roedores_filtered = df_roedores[df_roedores['Sede'] == sede].copy()
    df_lamparas_filtered = df_lamparas[df_lamparas['Sede'] == sede].copy()
    
    # Inicializar gestor de datos
    data_manager = ReportDataManager()
    
    # Crear documento
    doc = Document()
    
    # ===== HEADER DEL DOCUMENTO =====
    _add_document_header(doc, sede)
    
    # ===== SECCIÓN PREVENTIVOS =====
    print(f"\n📊 Procesando sección: SERVICIOS PREVENTIVOS")
    _process_preventivos_section(doc, df_preventivo_filtered, data_manager, sede)
    
    # ===== SECCIÓN ROEDORES =====
    print(f"\n📊 Procesando sección: CONTROL DE ROEDORES")
    _process_roedores_section(doc, df_roedores_filtered, data_manager, sede)
    
    # ===== SECCIÓN LÁMPARAS =====
    print(f"\n📊 Procesando sección: CONTROL DE INSECTOS VOLADORES")
    _process_lamparas_section(doc, df_lamparas_filtered, data_manager, sede)
    
    # ===== RESUMEN GENERAL =====
    print(f"\n📋 Generando resumen general del reporte")
    general_prompt_data = data_manager.generate_general_summary_prompt(sede)
    
    # Añadir sección de resumen general al documento
    doc.add_heading("RESUMEN EJECUTIVO GENERAL", level=1)
    doc.add_paragraph("Esta sección contiene el análisis integral de todo el reporte de control de plagas.")
    
    # Añadir prompt de resumen general
    if general_prompt_data:
        _add_prompt_to_doc(doc, general_prompt_data['prompt'], general_prompt_data['title'])
    
    # ===== EXPORTAR PROMPTS =====
    prompts_file = f"outputs/prompts_generados_{sede}.txt"
    data_manager.export_prompts_to_file(prompts_file)
    
    # Guardar documento
    output_path = f"outputs/reporte_serviplagas_{sede}.docx"
    doc.save(output_path)
    print(f"\n✅ Reporte guardado en: {output_path}")
    print(f"📝 Prompts disponibles en: {prompts_file}")
    
    # Retornar tanto la ruta del reporte como el gestor de datos para manipulación
    return output_path, data_manager


def _add_document_header(doc: Document, sede: str):
    """Añade el header del documento con logo y título."""
    # Agregar logo
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture('Logo/logo2021.png', width=Inches(3.25))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar título
    doc.add_paragraph(" ")
    doc.add_heading(f"Reporte Mensual {sede} - Serviplagas", level=1)
    doc.add_paragraph(" ")
    
    # ⚠️ ADVERTENCIA TEMPORAL SOBRE LOS PROMPTS (solo si están habilitados)
    if DOCUMENT_CONFIG.get('include_prompts_in_document', False):
        warning_para = doc.add_paragraph()
        warning_run = warning_para.add_run("⚠️ DOCUMENTO DE REVISIÓN - Los prompts de LLM mostrados en este documento son temporales para validación del sistema. En la versión final, estos prompts serán reemplazados por las respuestas generadas del LLM.")
        warning_run.font.bold = True
        warning_run.font.color.rgb = RGBColor(255, 102, 0)  # Naranja
        warning_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(" ")
    
    # Texto introductorio
    intro_text = (
        "A continuación, el resultado del ejercicio del séptimo mes (marzo) del ciclo de análisis de 12 meses "
        "con un nuevo recorrido mensual del PROGRAMA DE PREVENCIÓN DE PLAGAS DE IMPORTANCIA EN SALUD PÚBLICA (PPPISP) "
        "en las instalaciones del Hospital San Vicente Rionegro, se realizaron rondas preventivas del cronograma mensual "
        "con el propósito de controlar la presencia de insectos rastreros, voladores y roedores de menor importancia "
        "en los 4 bloques asistenciales y administrativos. Adicionalmente, se llevaron a cabo inspecciones en las "
        "estaciones de control de roedores y en las estaciones de luz para a combatir los insectos voladores. "
        "Asimismo, se efectuaron labores de mantenimiento en áreas verdes, espacios comunes, sumideros y sistemas "
        "de alcantarillado. Por último, se atendieron todas las novedades y mantenimientos correctivos solicitados "
        "a través del formato o por medio del grupo de WhatsApp."
    )
    doc.add_paragraph(intro_text)
    doc.add_paragraph(" ")


def _process_preventivos_section(doc: Document, df_preventivo: pd.DataFrame, 
                               data_manager: ReportDataManager, sede: str):
    """Procesa la sección de servicios preventivos."""
    
    # Preventivos 1: Órdenes vs Áreas
    doc.add_heading("Preventivos 1", level=2)
    doc.add_paragraph("La gráfica # 1 refleja la cantidad de órdenes de mantenimiento recibidas (con código), "
                     "la cantidad de áreas realizadas efectivamente y la cantidad de áreas con evidencia de plagas.")
    
    table_data, figure = generate_order_area_plot(df_preventivo)
    prompt_data = data_manager.generate_table_prompt('preventivos', 'order_area', table_data, sede)
    
    add_plot_to_doc(doc, figure)
    _add_table_to_doc(doc, table_data)
    
    # Añadir prompt para revisión
    _add_prompt_to_doc(doc, prompt_data['prompt'], prompt_data['title'])

    # Preventivos 2: Especies de plagas
    doc.add_heading("Preventivos 2", level=2)
    doc.add_paragraph("La gráfica # 2 refleja la relación por especie encontrada")
    
    table_data, figure = generate_plagas_timeseries_facet(df_preventivo)
    prompt_data = data_manager.generate_table_prompt('preventivos', 'plagas_species', table_data, sede)
    
    add_plot_to_doc(doc, figure)
    _add_table_to_doc(doc, table_data)
    
    # Añadir prompt para revisión
    _add_prompt_to_doc(doc, prompt_data['prompt'], prompt_data['title'])

    # Preventivos 3: Tendencia total
    doc.add_heading("Preventivos 3", level=2)
    doc.add_paragraph("La gráfica # 3 refleja la tendencia de eliminación mensual en preventivos")
    
    table_data, figure = generate_total_plagas_trend_plot(df_preventivo)
    prompt_data = data_manager.generate_table_prompt('preventivos', 'total_trend', table_data, sede)
    
    add_plot_to_doc(doc, figure)
    _add_table_to_doc(doc, table_data)
    
    # Añadir prompt para revisión
    _add_prompt_to_doc(doc, prompt_data['prompt'], prompt_data['title'])
    
    # Añadir prompt de resumen de sección
    section_prompt_data = data_manager.generate_section_summary_prompt('preventivos', sede)
    if section_prompt_data:
        _add_prompt_to_doc(doc, section_prompt_data['prompt'], section_prompt_data['title'])


def _process_roedores_section(doc: Document, df_roedores: pd.DataFrame, 
                            data_manager: ReportDataManager, sede: str):
    """Procesa la sección de control de roedores."""
    
    # Roedores 1: Estado de estaciones
    doc.add_heading("Roedores 1", level=2)
    doc.add_paragraph("La gráfica # 1 refleja la cantidad de HALLAZGOS Y NOVEDADES encontradas en las estaciones "
                     "portacebos instaladas en el hospital Universitario, dando cuenta de las tendencias en los puntos "
                     "de control y los puntos donde más se consume cebo rodenticida y su relacionamiento con la "
                     "disminución o la proliferación de esta especie en el tiempo; así como las estaciones en otros estados.")
    
    table_data, figure = generate_roedores_station_status_plot(df_roedores)
    prompt_data = data_manager.generate_table_prompt('roedores', 'station_status', table_data, sede)
    
    add_plot_to_doc(doc, figure)
    _add_table_to_doc(doc, table_data)
    
    # Añadir prompt para revisión
    _add_prompt_to_doc(doc, prompt_data['prompt'], prompt_data['title'])

    # Roedores 2: Tendencia de eliminación
    doc.add_heading("Roedores 2", level=2)
    doc.add_paragraph("La gráfica # 3 refleja la tendencia de consumo por mes comenzando el análisis "
                     "en el mes de septiembre de 2024")
    
    table_data, figure = plot_tendencia_eliminacion_mensual(df_roedores)
    prompt_data = data_manager.generate_table_prompt('roedores', 'elimination_trend', table_data, sede)
    
    add_plot_to_doc(doc, figure)
    _add_table_to_doc(doc, table_data)
    
    # Añadir prompt para revisión
    _add_prompt_to_doc(doc, prompt_data['prompt'], prompt_data['title'])
    
    # Añadir prompt de resumen de sección
    section_prompt_data = data_manager.generate_section_summary_prompt('roedores', sede)
    if section_prompt_data:
        _add_prompt_to_doc(doc, section_prompt_data['prompt'], section_prompt_data['title'])


def _process_lamparas_section(doc: Document, df_lamparas: pd.DataFrame, 
                            data_manager: ReportDataManager, sede: str):
    """Procesa la sección de control de lámparas."""
    
    # Lámparas 1: Estado por mes
    doc.add_heading("Lámparas 1", level=2)
    doc.add_paragraph("La gráfica # 1 refleja el consolidado del estado de las lámparas en el tiempo")
    
    table_data, figure = plot_estado_lamparas_por_mes(df_lamparas)
    prompt_data = data_manager.generate_table_prompt('lamparas', 'status_monthly', table_data, sede)
    
    add_plot_to_doc(doc, figure)
    _add_table_to_doc(doc, table_data)
    
    # Añadir prompt para revisión
    _add_prompt_to_doc(doc, prompt_data['prompt'], prompt_data['title'])

    # Lámparas 2: Estado con leyenda
    doc.add_heading("Lámparas 2", level=2)
    doc.add_paragraph("La gráfica # 2 refleja el estado de las lámparas por condición de la estación")
    
    table_data, figure = plot_estado_lamparas_con_leyenda(df_lamparas)
    prompt_data = data_manager.generate_table_prompt('lamparas', 'status_legend', table_data, sede)
    
    add_plot_to_doc(doc, figure)
    _add_table_to_doc(doc, table_data)
    
    # Añadir prompt para revisión
    _add_prompt_to_doc(doc, prompt_data['prompt'], prompt_data['title'])

    # Lámparas 3: Capturas por especie
    doc.add_heading("Lámparas 3", level=2)
    doc.add_paragraph("La gráfica # 3 refleja la cantidad de hallazgos por lámpara")
    
    table_data, figure = plot_capturas_especies_por_mes(df_lamparas)
    prompt_data = data_manager.generate_table_prompt('lamparas', 'captures_species', table_data, sede)
    
    add_plot_to_doc(doc, figure)
    _add_table_to_doc(doc, table_data)
    
    # Añadir prompt para revisión
    _add_prompt_to_doc(doc, prompt_data['prompt'], prompt_data['title'])

    # Lámparas 4: Tendencia de capturas
    doc.add_heading("Lámparas 4", level=2)
    doc.add_paragraph("La gráfica # 4 refleja el nivel de captura por mes")
    
    table_data, figure = plot_tendencia_total_capturas(df_lamparas)
    prompt_data = data_manager.generate_table_prompt('lamparas', 'captures_trend', table_data, sede)
    
    add_plot_to_doc(doc, figure)
    _add_table_to_doc(doc, table_data)
    
    # Añadir prompt para revisión
    _add_prompt_to_doc(doc, prompt_data['prompt'], prompt_data['title'])
    
    # Añadir prompt de resumen de sección
    section_prompt_data = data_manager.generate_section_summary_prompt('lamparas', sede)
    if section_prompt_data:
        _add_prompt_to_doc(doc, section_prompt_data['prompt'], section_prompt_data['title'])


def _add_table_to_doc(doc: Document, table_data: pd.DataFrame):
    """Añade una tabla de pandas al documento de Word."""
    if table_data.empty:
        return
        
    # Crear tabla en el documento
    table = doc.add_table(rows=1, cols=len(table_data.columns))
    table.style = 'Table Grid'
    
    # Añadir headers
    header_cells = table.rows[0].cells
    for i, column_name in enumerate(table_data.columns):
        header_cells[i].text = str(column_name)
    
    # Añadir filas de datos
    for _, row in table_data.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)


# Mantener función original para compatibilidad hacia atrás
def generate_report_in_memory(df_preventivo, df_roedores, df_lamparas, sede: str):
    """
    Función original mantenida para compatibilidad hacia atrás.
    DEPRECADA: Usar generate_enhanced_report() en su lugar.
    """
    print("⚠️  ADVERTENCIA: Usando función legacy. Considera migrar a generate_enhanced_report()")
    return _generate_legacy_report(df_preventivo, df_roedores, df_lamparas, sede)


def _generate_legacy_report(df_preventivo, df_roedores, df_lamparas, sede: str):
    df_preventivo = df_preventivo[df_preventivo['Sede'] == sede].copy()
    df_roedores = df_roedores[df_roedores['Sede'] == sede].copy()
    df_lamparas = df_lamparas[df_lamparas['Sede'] == sede].copy()   

    doc = Document()

    # Agregar logo
    # agregar el logo y obtener el parrafo en el que se insertó
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture('Logo/logo2021.png', width=Inches(3.25))
    # Center the paragraph
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar título
    doc.add_paragraph(" ")
    doc.add_heading(f"Reporte Mensual {sede} - Serviplagas", level=1)
    doc.add_paragraph(" ")
    doc.add_paragraph("A continuación, el resultado del ejercicio del séptimo mes (marzo) del ciclo de análisis de 12 meses con un nuevo recorrido mensual del PROGRAMA DE PREVENCIÓN DE PLAGAS DE IMPORTANCIA EN SALUD PÚBLICA (PPPISP) en las instalaciones del Hospital San Vicente Rionegro, se realizaron rondas preventivas del cronograma mensual con el propósito de controlar la presencia de insectos rastreros, voladores y roedores de menor importancia en los 4 bloques asistenciales y administrativos. Adicionalmente, se llevaron a cabo inspecciones en las estaciones de control de roedores y en las estaciones de luz para a combatir los insectos voladores. Asimismo, se efectuaron labores de mantenimiento en áreas verdes, espacios comunes, sumideros y sistemas de alcantarillado. Por último, se atendieron todas las novedades y mantenimientos correctivos solicitados a través del formato o por medio del grupo de WhatsApp.")
    doc.add_paragraph(" ")
    

    # Sección Preventivos
    doc.add_heading("Preventivos 1", level=2)
    doc.add_paragraph("La gráfica # 1 refleja la cantidad de órdenes de mantenimiento recibidas (con código), la cantidad de áreas realizadas efectivamente y la cantidad de áreas con evidencia de plagas.")
    table_1, fig1 = generate_order_area_plot(df_preventivo)
    add_plot_to_doc(doc, fig1)

    # Add table
    t = doc.add_table(rows=1, cols=len(table_1.columns))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(table_1.columns):
        hdr_cells[i].text = str(col)
    for _, row in table_1.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)


    doc.add_heading("Preventivos 2", level=2)
    doc.add_paragraph("La gráfica # 2 refleja la relación por especie encontrada")
    table_2, fig2= generate_plagas_timeseries_facet(df_preventivo)
    add_plot_to_doc(doc, fig2)

    # Add table
    t = doc.add_table(rows=1, cols=len(table_2.columns))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(table_2.columns):
        hdr_cells[i].text = str(col)
    for _, row in table_2.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)


    doc.add_heading("Preventivos 3", level=2)
    doc.add_paragraph("La gráfica # 3 refleja la tendencia de eliminación mensual en preventivos")
    table_3, fig3 = generate_total_plagas_trend_plot(df_preventivo)
    add_plot_to_doc(doc, fig3)

    # Add table
    t = doc.add_table(rows=1, cols=len(table_3.columns))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(table_3.columns):
        hdr_cells[i].text = str(col)
    for _, row in table_3.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)


    doc.add_heading("Roedores 1", level=2)
    doc.add_paragraph("La gráfica # 1 refleja la cantidad de HALLAZGOS Y NOVEDADES encontradas en las estaciones portacebos instaladas en el hospital Universitario, dando cuenta de las tendencias en los puntos de control y los puntos donde más se consume cebo rodenticida y su relacionamiento con la disminución o la proliferación de esta especie en el tiempo; así como las estaciones en otros estados.")
    table_1, fig1 = generate_roedores_station_status_plot(df_roedores)
    add_plot_to_doc(doc, fig1)

    # Add table
    t = doc.add_table(rows=1, cols=len(table_1.columns))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(table_1.columns):
        hdr_cells[i].text = str(col)
    for _, row in table_1.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)


    doc.add_heading("Roedores 2", level=2)
    doc.add_paragraph("LLa gráfica # 3 refleja la tendencia de consumo por mes comenzando el análisis en el mes de septiembre de 2024")
    table_2, fig2= plot_tendencia_eliminacion_mensual(df_roedores)
    add_plot_to_doc(doc, fig2)

        # Add table
    t = doc.add_table(rows=1, cols=len(table_2.columns))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(table_2.columns):
        hdr_cells[i].text = str(col)
    for _, row in table_2.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.add_heading("Lámparas 1", level=2)
    doc.add_paragraph("La gráfica # 1 refleja el consolidado del estado de las lámparas en el tiempo")
    table_1, fig2= plot_estado_lamparas_por_mes(df_lamparas)
    add_plot_to_doc(doc, fig2)

    # Add table
    t = doc.add_table(rows=1, cols=len(table_1.columns))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(table_1.columns):
        hdr_cells[i].text = str(col)
    for _, row in table_1.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)


    doc.add_heading("Lámparas 2", level=2)
    doc.add_paragraph("La gráfica # 2 refleja el estado de las lámparas por condición de la estación")
    table_2, fig2= plot_estado_lamparas_con_leyenda(df_lamparas)
    add_plot_to_doc(doc, fig2)

    # Add table
    t = doc.add_table(rows=1, cols=len(table_2.columns))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(table_2.columns):
        hdr_cells[i].text = str(col)
    for _, row in table_2.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)



    doc.add_heading("Lámparas 3", level=2)
    doc.add_paragraph("La gráfica # 3 refleja la cantidad de hallazgos por lámpara")
    table_3, fig2= plot_capturas_especies_por_mes(df_lamparas)
    add_plot_to_doc(doc, fig2)


        # Add table
    t = doc.add_table(rows=1, cols=len(table_3.columns))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(table_3.columns):
        hdr_cells[i].text = str(col)
    for _, row in table_3.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.add_heading("Lámparas 4", level=2)
    doc.add_paragraph("La gráfica # 4 refleja el nivel de captura por mes")
    table_4, fig2= plot_tendencia_total_capturas(df_lamparas)
    add_plot_to_doc(doc, fig2)

        # Add table
    t = doc.add_table(rows=1, cols=len(table_4.columns))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(table_4.columns):
        hdr_cells[i].text = str(col)
    for _, row in table_4.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)



    # Guardar el documento en disco
    output_path = f"outputs/reporte_serviplagas_{sede}.docx"
    doc.save(output_path)
    print(f" ✅  Reporte guardado en: {output_path}")